#!/usr/bin/env python3
# _*_ coding:utf-8 _*_

"""
读取Word，并把每行数据以N个空格分开成多个字段。最终把每行中的多个字段写入Excel。
"""

__author__ = 'Jans'
__date__ = '2018/2/22 上午10:19'
__createby__ = 'PyCharm'

import docx
import xlwt

# 获取文档对象
file = docx.Document("/Users/apple/Documents/test/111.docx")
print("段落数:" + str(len(file.paragraphs)))  # 每个回车隔离一段
# 输出每一段的内容
# for para in file.paragraphs:
#     print(para.text)
# 输出段落编号及段落内容
# for i in range(len(file.paragraphs)):
#     print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)


# 创建一个Workbook对象，这就相当于创建了一个Excel文件
book = xlwt.Workbook(encoding='utf-8', style_compression=0)
# 创建一个sheet对象，一个sheet对象对应Excel文件中的一张表格。cell_overwrite_ok，表示是否可以覆盖单元格，其实是Worksheet实例化的一个参数，默认值是False
sheet = book.add_sheet('test_sheet1', cell_overwrite_ok=True)

i = 0
for para in file.paragraphs:
    line = para.text
    columns = line.split()  # 默认为所有的空字符，包括空格、换行(\n)、制表符(\t)等。
    # print(columns)
    j = 0
    for col in columns:
        # print(col)
        # 此处需要将中文字符串解码成unicode码，否则会报错
        sheet.write(i, j, col)
        j = j + 1
    i = i + 1
book.save(r'/Users/apple/Documents/test/test1.xls')  # 该模块只能生成xls，不能生成xlsx。在字符串前加r，声明为raw字符串，这样就不会处理其中的转义了。否则，可能会报错
